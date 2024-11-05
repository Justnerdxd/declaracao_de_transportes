# modules/word_generator.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerar_word(output_word_path, dados):
    """
    Gera um documento Word com o layout conforme o modelo PDF fornecido.
    
    Parâmetros:
    - output_word_path: Caminho para salvar o documento Word preenchido.
    - dados: Dicionário com os dados a serem inseridos no documento.
    """
    # Cria um novo documento Word
    doc = Document()

    # Título
    titulo = doc.add_paragraph()
    run_titulo = titulo.add_run("DECLARAÇÃO DE TRANSPORTES")
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Empresa e CNPJ
    empresa = doc.add_paragraph()
    empresa_run = empresa.add_run("Screw Ind. Metalmecânica Eireli   CNPJ: 00.397.908/0001-80")
    empresa_run.font.size = Pt(12)
    empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adicionar espaço após o título
    doc.add_paragraph().add_run("\n")

    # Sessão de informações do motorista e da coleta
    doc.add_paragraph(f"Motorista: {dados['Motorista']}    CPF/RG: {dados['CPF_RG']}    Telefone: {dados['Telefone']}")
    doc.add_paragraph(f"Cidade: {dados['Cidade']}    Código: __________")
    doc.add_paragraph(f"Placa: {dados['Placa']}    Transportadora: {dados['Transportadora']}")
    doc.add_paragraph(f"CNPJ: {dados['CNPJ_Transportadora']}    Data: {dados['Data']}    Hora: {dados['Hora']}")

    # Destinatário e Expedição
    doc.add_paragraph().add_run("\nDestinatário:")
    doc.add_paragraph().add_run("\nExpedição")
    doc.add_paragraph(f"Nº Matrícula: __________  Nº Matrícula: __________  Doc.Ident: __________")

    # Declaração
    declaracao = doc.add_paragraph()
    declaracao_run = declaracao.add_run(
        "\nDeclaro, para fins de comprovação de ambas as partes, que estou ciente e de acordo com as informações "
        "que constam neste documento e em nota fiscal nele relacionadas, conferidos e confirmados no momento do recebimento."
    )
    declaracao_run.font.size = Pt(10)

    # Faturamento e Motorista/Destinatário
    doc.add_paragraph().add_run("\nFaturamento  Motorista/Destinatário\n")
    doc.add_paragraph("Outros")

    # Critérios de Avaliação
    doc.add_paragraph("\nAcondicionamento da carga satisfatório? [  ] SIM   [  ] NÃO")
    doc.add_paragraph("Condições do veículo satisfatória? [  ] SIM   [  ] NÃO")

    # Tabela para Notas Fiscais
    doc.add_paragraph("\nCritérios de Avaliação Geral\n")
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nº Nota Fiscal'
    hdr_cells[1].text = 'Peso Descrito'
    hdr_cells[2].text = 'Nº Volumes'
    hdr_cells[3].text = 'Observações'

    # Adiciona as notas fiscais na tabela
    for nota in dados['Notas_Fiscais']:
        row_cells = table.add_row().cells
        row_cells[0].text = nota['numero']
        row_cells[1].text = nota['peso']
        row_cells[2].text = nota['volumes']
        row_cells[3].text = nota['observacao']

    # Salva o documento Word no caminho especificado
    doc.save(output_word_path)
    print(f"Documento Word gerado e salvo em: {output_word_path}")
